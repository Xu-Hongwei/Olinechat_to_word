from __future__ import annotations

import argparse
import json
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


BODY_CN_FONT = "宋体"
BODY_EN_FONT = "Times New Roman"
CODE_FONT = "Consolas"
HEADING_SIZES = {1: 22, 2: 16, 3: 15, 4: 14, 5: 14, 6: 12}


def ensure_rfonts(target, ascii_font: str, east_asia_font: str) -> None:
    r_pr = target._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), ascii_font)
    r_fonts.set(qn("w:hAnsi"), ascii_font)
    r_fonts.set(qn("w:cs"), ascii_font)
    r_fonts.set(qn("w:eastAsia"), east_asia_font)


def apply_run_font(run, *, cn_font: str = BODY_CN_FONT, en_font: str = BODY_EN_FONT, size: int = 12, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = en_font
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    ensure_rfonts(run, en_font, cn_font)


def add_paragraph_border(paragraph, border: str, color: str, size: str = "8") -> None:
    p_pr = paragraph._element.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    border_element = p_bdr.find(qn(f"w:{border}"))
    if border_element is None:
        border_element = OxmlElement(f"w:{border}")
        p_bdr.append(border_element)
    border_element.set(qn("w:val"), "single")
    border_element.set(qn("w:sz"), size)
    border_element.set(qn("w:space"), "1")
    border_element.set(qn("w:color"), color)


def add_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._element.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = BODY_EN_FONT
    normal.font.size = Pt(12)
    ensure_rfonts(normal, BODY_EN_FONT, BODY_CN_FONT)
    normal.paragraph_format.first_line_indent = Cm(0.74)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    for level, size in HEADING_SIZES.items():
        style = document.styles[f"Heading {level}"]
        style.font.name = BODY_EN_FONT
        style.font.size = Pt(size)
        style.font.bold = True
        ensure_rfonts(style, BODY_EN_FONT, BODY_CN_FONT)
        style.paragraph_format.first_line_indent = Cm(0)
        style.paragraph_format.line_spacing = 1.2
        style.paragraph_format.space_before = Pt(12 if level <= 2 else 6)
        style.paragraph_format.space_after = Pt(6)


def apply_body_paragraph_format(paragraph, *, first_line_indent: bool = True, left_indent_cm: float = 0.0, alignment: int | None = None) -> None:
    paragraph.paragraph_format.left_indent = Cm(left_indent_cm)
    paragraph.paragraph_format.right_indent = Cm(0)
    paragraph.paragraph_format.first_line_indent = Cm(0.74) if first_line_indent else Cm(0)
    paragraph.paragraph_format.hanging_indent = Cm(0)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    if alignment is not None:
        paragraph.alignment = alignment


def add_link_run(paragraph, text: str, url: str, *, size: int = 12) -> None:
    display_text = text.strip() if text.strip() else url
    link_run = paragraph.add_run(display_text)
    apply_run_font(link_run, size=size)
    link_run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
    link_run.underline = True
    if display_text != url:
        tail_run = paragraph.add_run(f" ({url})")
        apply_run_font(tail_run, size=size)


def resolve_image_path(source: str, source_dir: str | None) -> Path | None:
    source = source.strip()
    if not source:
        return None
    if re.match(r"^[A-Za-z][A-Za-z0-9+.-]*://", source):
        return None
    candidate = Path(source)
    if candidate.is_absolute():
        return candidate if candidate.exists() else None
    if source_dir:
        relative = Path(source_dir) / candidate
        if relative.exists():
            return relative
    if candidate.exists():
        return candidate
    return None

class MathCursor:
    def __init__(self, math_items: list[dict]) -> None:
        self.math_items = math_items
        self.index = 0

    def next_placeholder(self) -> str:
        if self.index >= len(self.math_items):
            placeholder = f"[[EQ_MISSING_{self.index}]]"
            self.index += 1
            return placeholder
        placeholder = self.math_items[self.index]["placeholder"]
        self.index += 1
        return placeholder


def add_inline_runs(
    paragraph,
    runs: list[dict],
    math_cursor: MathCursor,
    heading_level: int | None = None,
    source_dir: str | None = None,
) -> None:
    size = HEADING_SIZES.get(heading_level, 12) if heading_level else 12
    for run in runs:
        run_type = run["type"]
        if run_type == "math":
            text_run = paragraph.add_run(math_cursor.next_placeholder())
            apply_run_font(text_run, cn_font="Cambria Math", en_font="Cambria Math", size=size)
            continue
        if run_type == "link":
            add_link_run(paragraph, run.get("text", run.get("url", "")), run.get("url", ""), size=size)
            continue
        if run_type == "image":
            source = run.get("src", "")
            alt_text = run.get("alt", "image")
            resolved_path = resolve_image_path(source, source_dir)
            if resolved_path:
                image_run = paragraph.add_run()
                image_run.add_picture(str(resolved_path), width=Cm(12))
            else:
                fallback = paragraph.add_run(f"[图片: {alt_text}]({source})")
                apply_run_font(fallback, size=size, italic=True)
            continue

        text = run.get("text", "")
        if not text:
            continue
        text_run = paragraph.add_run(text)
        if run_type == "bold":
            apply_run_font(text_run, size=size, bold=True)
        elif run_type == "italic":
            apply_run_font(text_run, size=size, italic=True)
        elif run_type == "code":
            apply_run_font(text_run, cn_font=CODE_FONT, en_font=CODE_FONT, size=11)
            text_run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        else:
            apply_run_font(text_run, size=size)


def add_list_paragraph(document: Document, item: dict, math_cursor: MathCursor, source_dir: str | None = None) -> None:
    paragraph = document.add_paragraph()
    apply_body_paragraph_format(paragraph, first_line_indent=False, left_indent_cm=0.74 * (item.get("level", 0) + 1))
    paragraph.paragraph_format.hanging_indent = Cm(0.6)

    if item.get("ordered"):
        prefix = f"{item.get('start', 1)}. "
    elif item.get("task") is True:
        prefix = "☑ "
    elif item.get("task") is False:
        prefix = "☐ "
    else:
        prefix = "- "
    prefix_run = paragraph.add_run(prefix)
    apply_run_font(prefix_run, size=12)
    add_inline_runs(paragraph, item["runs"], math_cursor, source_dir=source_dir)


def add_table_block(document: Document, block: dict, math_cursor: MathCursor, source_dir: str | None = None) -> None:
    header = block.get("header", [])
    rows = block.get("rows", [])
    column_count = max([len(header), *[len(row) for row in rows], 1])
    table = document.add_table(rows=1 + len(rows), cols=column_count)
    try:
        table.style = "Table Grid"
    except KeyError:
        pass

    for column_index in range(column_count):
        cell = table.rows[0].cells[column_index]
        cell_runs = header[column_index] if column_index < len(header) else []
        paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
        apply_body_paragraph_format(paragraph, first_line_indent=False)
        add_inline_runs(paragraph, cell_runs, math_cursor, source_dir=source_dir)
        for run in paragraph.runs:
            run.bold = True

    for row_index, row in enumerate(rows, start=1):
        for column_index in range(column_count):
            cell = table.rows[row_index].cells[column_index]
            cell_runs = row[column_index] if column_index < len(row) else []
            paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
            apply_body_paragraph_format(paragraph, first_line_indent=False)
            add_inline_runs(paragraph, cell_runs, math_cursor, source_dir=source_dir)


def build_document(payload: dict, output_docx: str) -> None:
    document = Document()
    configure_styles(document)
    math_cursor = MathCursor(payload.get("math_items", []))
    source_dir = payload.get("source_dir")
    first_heading_level_one = True

    for block in payload.get("blocks", []):
        block_type = block["type"]
        if block_type == "heading":
            level = max(1, min(block.get("level", 1), 6))
            if level == 1 and not first_heading_level_one:
                document.add_page_break()
            paragraph = document.add_paragraph(style=f"Heading {level}")
            paragraph.paragraph_format.keep_with_next = True
            if level == 1:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                first_heading_level_one = False
            add_inline_runs(paragraph, block["runs"], math_cursor, heading_level=level, source_dir=source_dir)
        elif block_type == "paragraph":
            paragraph = document.add_paragraph(style="Normal")
            apply_body_paragraph_format(paragraph)
            add_inline_runs(paragraph, block["runs"], math_cursor, source_dir=source_dir)
        elif block_type == "blockquote":
            paragraph = document.add_paragraph(style="Normal")
            apply_body_paragraph_format(paragraph, first_line_indent=False, left_indent_cm=1.0)
            add_paragraph_border(paragraph, "left", "6366F1", "12")
            add_inline_runs(paragraph, block["runs"], math_cursor, source_dir=source_dir)
            for run in paragraph.runs:
                run.italic = True
        elif block_type == "code_block":
            paragraph = document.add_paragraph()
            apply_body_paragraph_format(paragraph, first_line_indent=False, left_indent_cm=0.5)
            add_paragraph_shading(paragraph, "F5F5F5")
            if block.get("language"):
                label_run = paragraph.add_run(f"[{block['language']}]\n")
                apply_run_font(label_run, cn_font=CODE_FONT, en_font=CODE_FONT, size=9, italic=True)
            code_run = paragraph.add_run(block.get("text", ""))
            apply_run_font(code_run, cn_font=CODE_FONT, en_font=CODE_FONT, size=9)
        elif block_type == "list":
            for item in block.get("items", []):
                add_list_paragraph(document, item, math_cursor, source_dir=source_dir)
        elif block_type == "table":
            add_table_block(document, block, math_cursor, source_dir=source_dir)
        elif block_type == "separator":
            paragraph = document.add_paragraph()
            apply_body_paragraph_format(paragraph, first_line_indent=False)
            add_paragraph_border(paragraph, "bottom", "CCCCCC", "6")
        elif block_type == "math_block":
            paragraph = document.add_paragraph()
            apply_body_paragraph_format(paragraph, first_line_indent=False, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            run = paragraph.add_run(math_cursor.next_placeholder())
            apply_run_font(run, cn_font="Cambria Math", en_font="Cambria Math", size=12)

    document.save(output_docx)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--input-json", required=True)
    parser.add_argument("--output-docx", required=True)
    args = parser.parse_args()

    with open(args.input_json, "r", encoding="utf-8") as handle:
        payload = json.load(handle)

    build_document(payload, args.output_docx)
    print(f"布局文档已生成：{args.output_docx}")


if __name__ == "__main__":
    main()
